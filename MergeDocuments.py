import os
import docx

# Define the folder containing the DOCX files and the output file path
folder_path = '/Users/annawilliams/Desktop/MergeWordDocuments'
output_file = '/Users/annawilliams/Desktop/MergeWordDocuments/merged_document.docx'

# Generate filenames (e.g., 'sample1.docx')
filenames = os.listdir(folder_path)

# Generate filepaths (e.g., 'word_documents/sample1.docx')
filepaths = [
    os.path.join(folder_path, filename)
    for filename in filenames
    if filename.endswith('.docx') and not filename.startswith('~$')  # Skip temporary files
]

# Create a new Document object for the merged document
merged_document = docx.Document()

# Iterate over all DOCX files in the specified folder
for filepath in filepaths:
    try:
        doc = docx.Document(filepath)

        # Append each paragraph from the current document to the merged document
        for para in doc.paragraphs:
            # Add each paragraph text to the merged document
            merged_document.add_paragraph(para.text)

    except Exception as e:
        print(f"Error processing {filepath}: {e}")

# Save the merged document to the specified output file
merged_document.save(output_file)